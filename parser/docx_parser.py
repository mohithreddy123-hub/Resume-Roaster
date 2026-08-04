"""
parser/docx_parser.py
---------------------
Extracts raw text from DOCX files using python-docx.
Handles: corrupted DOCX files, empty documents.
Extracts text from both paragraphs and tables.
Returns raw text string — no cleaning done here.
"""

import io
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DOCXParseError(Exception):
    """Raised when a DOCX cannot be parsed for any reason."""
    pass


class BlankDOCXError(DOCXParseError):
    """Raised when a DOCX has no extractable text."""
    pass


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract raw text from a DOCX file given its bytes.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Raw extracted text as a single string.

    Raises:
        DOCXParseError: If the file is corrupted or cannot be read.
        BlankDOCXError: If no text could be extracted.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError:
        raise DOCXParseError(
            "This file seems corrupted. Try re-saving it as a DOCX and upload again."
        )
    except Exception as e:
        raise DOCXParseError(
            "Could not open this DOCX file. It may be damaged or in an unsupported format."
        ) from e

    text_parts: list[str] = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    # Extract text from tables (important — resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    if not text_parts:
        raise BlankDOCXError(
            "This document appears to be empty. Please check the file and try again."
        )

    raw_text = "\n".join(text_parts)
    return raw_text
